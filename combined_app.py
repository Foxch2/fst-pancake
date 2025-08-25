# combined_app.py

import pandas as pd
import pyodbc
import configparser
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext
import os
from pathlib import Path
# Импорты для DOCX
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
# Импорты для QR-кода
import qrcode
from PIL import Image
from datetime import datetime, timedelta
import logging
import unicodedata
import re
import requests
import json
import base64
import uuid
import pytz
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
from io import BytesIO
import pytz # Убедитесь, что pytz импортирован в начале файла
                        

# --- Настройка логирования ---
LOG_FILE = "combined_app.log"
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(funcName)s:%(lineno)d - %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# --- Функция очистки текста для PDF/DOCX ---
def clean_text_for_pdf(text):
    if not isinstance(text, str):
        return text
    # Удаляем невидимые символы (control characters)
    visible_chars = ''.join(char for char in text if unicodedata.category(char)[0] != 'C')
    # Приводим к NFC (каноническая форма Unicode)
    cleaned = unicodedata.normalize('NFC', visible_chars)
    # Убираем лишние пробелы
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned



# --- Загрузка конфигурации ---
def load_config(config_file="conf.ini"):
    """
    Загружает конфигурацию из INI-файла.
    Возвращает словарь с настройками.
    """
    if not os.path.exists(config_file):
        error_msg = f"Файл конфигурации {config_file} не найден."
        logging.error(error_msg)
        raise FileNotFoundError(error_msg)

    config = configparser.ConfigParser()
    # config.optionxform = str # Оставляем стандартное поведение (ключи в нижнем регистре)

    try:
        config.read(config_file, encoding='utf-8')
    except Exception as e:
        error_msg = f"Ошибка чтения файла конфигурации {config_file}: {e}"
        logging.error(error_msg)
        raise

    # --- Проверка наличия обязательных секций ---
    required_sections = ['DATABASE', 'YOOCASSA', 'MAIL']
    for section in required_sections:
        if section not in config:
            error_msg = f"Секция [{section}] не найдена в {config_file}."
            logging.error(error_msg)
            raise KeyError(error_msg)

    # --- Безопасное извлечение и обработка секций ---
    try:
        # 1. DATABASE
        db_section = config['DATABASE']
        # Функция для безопасного получения строкового значения
        def get_str_option(section, key, default=''):
            value = section.get(key, default)
            # configparser обычно возвращает строки, но на всякий случай
            return str(value).strip() if value is not None else default

        db_config = {
            'server': get_str_option(db_section, 'server'),
            'database': get_str_option(db_section, 'database'),
            'driver': get_str_option(db_section, 'driver'),
            # Безопасная обработка auth_mode
            'auth_mode': get_str_option(db_section, 'auth_mode').strip().lower(),
            'username': get_str_option(db_section, 'username'),
            'password': get_str_option(db_section, 'password'),
        }

        # 2. YOOCASSA
        yookassa_section = config['YOOCASSA']
        yookassa_config = {
            'shop_id': get_str_option(yookassa_section, 'shop_id').strip(),
            'secret_key': get_str_option(yookassa_section, 'secret_key').strip(),
            # invoice_lifetime_hours больше не используется, но можно оставить для совместимости
            # 'invoice_lifetime_hours': int(yookassa_section.get('invoice_lifetime_hours', 168)),
            'vat_code': int(yookassa_section.get('vat_code', 4)),
            'tax_system_code': int(yookassa_section.get('tax_system_code', 1)),
        }

        # 3. MAIL
        mail_section = config['MAIL']
        mail_config = {
            'smtp_server': get_str_option(mail_section, 'smtp_server').strip(),
            'smtp_port': int(mail_section.get('smtp_port', 25)),
            'use_tls': get_str_option(mail_section, 'use_tls').strip().lower() in ['true', '1', 'yes', 'on'],
            'sender_email': get_str_option(mail_section, 'sender_email').strip(),
            'sender_password': get_str_option(mail_section, 'sender_password').strip(),
            'sender_name': get_str_option(mail_section, 'sender_name', 'No Reply').strip(),
        }

        # 4. OPTIONS (необязательная секция)
        options_config = {'encoding': 'utf-8', 'MAXWeight': 0} # Значения по умолчанию
        if 'OPTIONS' in config:
            options_section = config['OPTIONS']
            options_config['encoding'] = get_str_option(options_section, 'encoding', 'utf-8')
            # Обработка MAXWeight
            max_weight_str = get_str_option(options_section, 'MAXWeight', '0')
            try:
                options_config['MAXWeight'] = int(max_weight_str)
                if options_config['MAXWeight'] < 0:
                     logging.warning(f"Значение MAXWeight в [OPTIONS] отрицательное ({options_config['MAXWeight']}), установлено в 0.")
                     options_config['MAXWeight'] = 0
            except ValueError:
                logging.warning(f"Некорректное значение MAXWeight='{max_weight_str}' в [OPTIONS], использую 0 (без ограничений).")
                options_config['MAXWeight'] = 0
        else:
            logging.info("Секция [OPTIONS] не найдена в конфигурации, используются значения по умолчанию.")

    except Exception as processing_error:
        error_msg = f"Ошибка обработки секций конфигурации: {processing_error}"
        logging.error(error_msg, exc_info=True) # exc_info=True для полной трассировки
        raise # Повторно вызываем исключение

    logging.info("Конфигурация успешно загружена.")
    return {
        'DATABASE': db_config,
        'YOOCASSA': yookassa_config,
        'MAIL': mail_config,
        'OPTIONS': options_config
    }



# --- Подключение к БД ---
def connect_to_db(db_config):
    try:
        conn_str = (
            f"DRIVER={{{db_config['driver']}}};"
            f"SERVER={db_config['server']};"
            f"DATABASE={db_config['database']};"
        )
        if db_config['auth_mode'] == 'windows':
            conn_str += "Trusted_Connection=yes;"
        else:
            conn_str += f"UID={db_config['username']};PWD={db_config['password']};"

        # Подключение
        conn = pyodbc.connect(conn_str, autocommit=False)
        logging.info("Подключение к БД установлено.")
        return conn
    except Exception as e:
        error_msg = f"Ошибка подключения: {e}"
        logging.error(error_msg)
        return None

# --- Класс основного приложения ---
class CombinedApp:
    def __init__(self, root):
        logging.info("Инициализация комбинированного приложения...")
        self.root = root
        self.root.title("📦 Система заказов и счетов")
        self.root.geometry("800x900")
        self.root.resizable(True, True)

        self.conn = None
        self.csv_filename = "заказ"
        self.config = None

        # --- Загрузка конфигурации ---
        try:
            self.config = load_config()
        except Exception as e:
            logging.critical(f"Ошибка загрузки конфигурации: {e}")
            messagebox.showerror("❌ Критическая ошибка", f"Ошибка загрузки конфигурации:\n{e}")
            return

        # --- Автоматическое подключение к БД при запуске ---
        logging.info("Автоматическая попытка подключения к БД при запуске...")
        self.conn = connect_to_db(self.config['DATABASE'])
        if not self.conn:
            logging.critical("Не удалось подключиться к БД при запуске.")
            # messagebox.showerror("❌ Ошибка", "Не удалось подключиться к базе данных. Проверьте настройки в conf.ini.")

        # --- Создание вкладок ---
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Вкладка заказов
        self.orders_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.orders_frame, text="📦 Заказы")
        self.create_orders_widgets()

        # Вкладка счетов
        self.bills_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.bills_frame, text="💳 Счета")
        self.create_bills_widgets()

        # --- Инициализация GUI ---
        self.init_gui_state()

        logging.info("Инициализация комбинированного приложения завершена.")

    def init_gui_state(self):
        """Инициализация состояния элементов GUI"""
        if self.conn:
            # Активируем кнопки на вкладке заказов
            self.btn_load.config(state=tk.NORMAL)
            self.btn_print.config(state=tk.NORMAL)
            self.load_existing_loads()
            
            # Активируем кнопки на вкладке счетов
            self.btn_get_info.config(state=tk.NORMAL)
            yookassa_config = self.config.get('YOOCASSA', {})
            shop_id = yookassa_config.get('shop_id', '')
            secret_key = yookassa_config.get('secret_key', '')
            self.btn_auto_create.config(
                state=tk.NORMAL if shop_id and secret_key else tk.DISABLED
            )
        else:
            # Деактивируем кнопки на обеих вкладках
            self.btn_load.config(state=tk.DISABLED)
            self.btn_print.config(state=tk.DISABLED)
            self.btn_get_info.config(state=tk.DISABLED, text="🚫 Нет подключения к БД")
            self.btn_auto_create.config(state=tk.DISABLED, text="🚫 Нет подключения к БД")
            
        # Проверка статуса API ЮKassa
        yookassa_config = self.config.get('YOOCASSA', {})
        shop_id = yookassa_config.get('shop_id', '')
        secret_key = yookassa_config.get('secret_key', '')
        if not shop_id or not secret_key:
            self.api_status_label.config(text="❌ API не настроен", foreground="red")
        else:
            self.api_status_label.config(text="✅ API настроен", foreground="green")

    # ==================== ФУНКЦИИ ДЛЯ ЗАКАЗОВ ====================
    
    def create_orders_widgets(self):
        """Создание элементов для вкладки заказов"""
        container = ttk.Frame(self.orders_frame)
        container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        self.label = tk.Label(container, text="Загрузка и печать заказов", font=("Arial", 16))
        self.label.pack(pady=10)

        self.btn_load = tk.Button(container, text="📂 Загрузить CSV", command=self.load_csv, width=40, height=2, state=tk.DISABLED)
        self.btn_load.pack(pady=5)

        tk.Label(container, text="Выберите загрузку для печати:", font=("Arial", 12)).pack(pady=5)
        self.combo_loads = ttk.Combobox(container, width=60, state="readonly", height=20)
        self.combo_loads.pack(pady=5)

        self.btn_print = tk.Button(
            container,
            text="🖨️ Распечатать выбранную загрузку",
            command=self.print_selected_load,
            width=40,
            height=2,
            state=tk.DISABLED
        )
        self.btn_print.pack(pady=5)

        # Отображение статуса подключения
        if self.conn:
            tk.Label(container, text="✅ Подключение к БД установлено", fg="green").pack(pady=5)
        else:
            tk.Label(container, text="❌ Нет подключения к БД", fg="red").pack(pady=5)

    def load_existing_loads(self):
        """Загружает список уникальных загрузок из БД"""
        if not self.conn:
            return
        try:
            query = """
                SELECT DISTINCT LoadedID, LoadedName
                FROM Orders
                WHERE LoadedID IS NOT NULL
                ORDER BY LoadedID DESC
            """
            df = pd.read_sql(query, self.conn)
            if not df.empty:
                loads = [f"{row['LoadedID']} - {row['LoadedName']}" for _, row in df.iterrows()]
                self.combo_loads['values'] = loads
                if loads:
                    self.combo_loads.current(0)
            else:
                self.combo_loads['values'] = []
        except Exception as e:
            logging.error(f"Ошибка загрузки списка загрузок: {e}")

    def load_csv(self):
        if not self.conn:
            messagebox.showwarning("⚠️", "Сначала подключитесь к БД.")
            return

        file_path = filedialog.askopenfilename(
            title="Выберите CSV-файл",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        if not file_path:
            return

        self.csv_filename = Path(file_path).stem
        loaded_id = datetime.now().strftime("%d.%m.%Y-%H.%M.%S")
        logging.info(f"Загрузка: ID={loaded_id}, файл={file_path}")

        try:
            df = pd.read_csv(file_path, dtype=str, encoding='utf-8', on_bad_lines='skip', header=0)
            df.columns = df.columns.str.strip()

            # Переименование
            column_mapping = {
                "Номер заказа": "OrderID",
                "Имя": "Name",
                "email": "Email",
                "Табельный номер": "EmployeeID",
                "Телефон": "Phone",
                "Доставка": "DeliveryPoint",
                "GRD код": "GRDCode",
                "Наименование": "ProductName",
                "Вес кейса, Брутто/г": "CaseWeight",
                "Цена": "Price",
                "Количество": "Quantity",
                "Общий вес, Брутто/г": "TotalWeight",
                "Сумма": "TotalAmount"
            }
            df = df.rename(columns=column_mapping)

            # Приведение типов
            numeric_cols = ['OrderID', 'EmployeeID', 'CaseWeight', 'Price', 'Quantity', 'TotalWeight', 'TotalAmount']
            for col in numeric_cols:
                if col in df.columns:
                    df[col] = pd.to_numeric(df[col], errors='coerce')
            df = df.where(pd.notnull(df), None)

            # Удаление строки "ИТОГО"
            if 'OrderID' in df.columns:
                df = df.dropna(subset=["OrderID"])
                df = df[df["OrderID"] != "ИТОГО"]

            # === ПРОВЕРКА ДУБЛИКАТОВ: Только по существующим OrderID в БД ===
            cursor_check = self.conn.cursor()
            cursor_check.execute("SELECT OrderID FROM Orders")
            # Создаем множество существующих OrderID из БД
            existing_order_ids = {row[0] for row in cursor_check.fetchall() if row[0] is not None}
            logging.debug(f"Найдено {len(existing_order_ids)} существующих заказов в БД.")

            # Создаем множество OrderID из загружаемого файла
            # df['OrderID'] уже числовой благодаря pd.to_numeric
            file_order_ids = set(df['OrderID'].dropna().astype(int))
            logging.debug(f"Найдено {len(file_order_ids)} заказов в файле.")

            # Находим пересечение - дубликаты
            duplicate_order_ids = file_order_ids & existing_order_ids
            new_order_ids = file_order_ids - duplicate_order_ids

            if duplicate_order_ids:
                duplicates = [f"заказ № {order_id} был загружен ранее" for order_id in sorted(duplicate_order_ids)]
                msg = "Следующие заказы уже были загружены:\n" + "\n".join(duplicates)
                messagebox.showinfo("ℹ️ Повторы", msg)
                logging.info(f"Найдено {len(duplicate_order_ids)} дубликатов.")

            if not new_order_ids:
                messagebox.showinfo("ℹ️", "Нет новых заказов для загрузки.")
                return

            # Фильтруем DataFrame, оставляя только новые заказы
            new_orders = df[df['OrderID'].isin(new_order_ids)]
            logging.info(f"Выбрано {len(new_orders)} строк для новых заказов ({len(new_order_ids)} уникальных заказов).")

            # --- ВСТАВКА НОВЫХ ЗАКАЗОВ ---
            cursor = self.conn.cursor()
            inserted = 0
            for _, row in new_orders.iterrows():
                # Очищаем текстовые поля перед вставкой
                clean_row = [
                    clean_text_for_pdf(val) if isinstance(val, str) else val for val in row
                ]
                cursor.execute("""
                    INSERT INTO Orders (OrderID, Name, Email, EmployeeID, Phone, DeliveryPoint,
                                        GRDCode, ProductName, CaseWeight, Price, Quantity, TotalWeight, TotalAmount,
                                        LoadedID, LoadedName)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, tuple(clean_row) + (loaded_id, self.csv_filename))
                inserted += 1
            self.conn.commit()

            # === НАЧАЛО: Загрузка агрегированных данных в таблицу bills ===
            try:
                # Агрегируем данные по новым заказам из new_orders DataFrame
                # Группируем по OrderID, Name, Email и суммируем Price * Quantity для каждой строки
                aggregated_data = new_orders.groupby(['OrderID', 'Name', 'Email']).apply(
                    lambda group: pd.Series({
                        'total_amount': (group['Price'] * group['Quantity']).sum()
                    })
                ).reset_index()

                # Добавляем loaded_id ко всем строкам
                aggregated_data['loaded_id'] = loaded_id

                # Подготавливаем данные для вставки
                # Выбираем нужные столбцы и преобразуем в список кортежей
                bill_records = aggregated_data[
                    ['OrderID', 'Name', 'Email', 'total_amount', 'loaded_id']
                ].apply(lambda row: (
                    int(row['OrderID']),
                    clean_text_for_pdf(row['Name']),
                    clean_text_for_pdf(row['Email']),
                    float(row['total_amount']),
                    row['loaded_id']
                ), axis=1).tolist()

                # Вставка данных в таблицу bills
                cursor_bills = self.conn.cursor()
                insert_query = """
                    INSERT INTO bills (order_id_int, name, email, total_amount, loaded_id, status)
                    VALUES (?, ?, ?, ?, ?, 'pending')
                """
                if bill_records: # Проверяем, есть ли данные для вставки
                    cursor_bills.executemany(insert_query, bill_records)
                    self.conn.commit()
                    logging.info(f"В таблицу 'bills' добавлено {len(bill_records)} записей.")
                else:
                    logging.info("Нет данных для добавления в таблицу 'bills'.")
            except Exception as e:
                error_msg = f"Ошибка при добавлении данных в таблицу 'bills': {e}"
                logging.error(error_msg)
                # Можно выбрать, показывать ли ошибку пользователю или продолжить
                # messagebox.showerror("❌ Ошибка", error_msg)
                # Для надежности лучше логировать и продолжать основной поток
            # === КОНЕЦ: Загрузка агрегированных данных в таблицу bills ===


            messagebox.showinfo("✅", f"Загружено {inserted} новых заказов.\nID: {loaded_id}")
            logging.info(f"Загружено {inserted} заказов с LoadedID={loaded_id}")

            # Обновляем список
            self.load_existing_loads()

        except Exception as e:
            error_msg = f"Ошибка загрузки CSV: {e}"
            logging.error(error_msg)
            messagebox.showerror("❌ Ошибка", error_msg)

    def print_selected_load(self):
        """Печать DOCX для выбранной загрузки, только оплаченные заказы."""
        selection = self.combo_loads.get()
        if not selection:
            messagebox.showwarning("⚠️", "Выберите загрузку для печати.")
            return

        # Извлекаем LoadedID из строки выбора, например, "13.08.2025-10.00.00 - Orders_2025-08-11"
        parts = selection.split(" - ", 1) # Разделяем максимум на две части
        if not parts:
            messagebox.showerror("❌ Ошибка", "Некорректный формат выбранной загрузки.")
            return

        loaded_id = parts[0]
        loaded_name = parts[1] if len(parts) > 1 else "Неизвестно"
        logging.info(f"Печать загрузки: ID={loaded_id}, имя={loaded_name}")

        try:
            # === ШАГ 1: Получаем список ОПЛАЧЕННЫХ заказов (order_id_int) из таблицы `bills` ===
            query_paid_bills = """
                SELECT order_id_int
                FROM bills
                WHERE loaded_id = ? AND status = 'succeeded'
            """
            # Используем read_sql с параметрами для безопасности
            df_paid_bills = pd.read_sql(query_paid_bills, self.conn, params=[loaded_id])

            # Проверяем, есть ли оплаченные заказы
            if df_paid_bills.empty:
                messagebox.showinfo("ℹ️", "В выбранной загрузке нет заказов со статусом 'succeeded'. Печать невозможна.")
                logging.info(f"Нет оплаченных заказов для загрузки {loaded_id}.")
                return

            # Создаем кортеж order_id_int для следующего запроса
            # Очень важно привести к int и убедиться, что нет None
            order_ids_paid_set = set(df_paid_bills['order_id_int'].dropna().astype(int))
            if not order_ids_paid_set:
                 messagebox.showinfo("ℹ️", "Не найдены корректные ID оплаченных заказов.")
                 logging.warning(f"Не найдены корректные ID оплаченных заказов для загрузки {loaded_id}.")
                 return

            order_ids_paid_tuple = tuple(order_ids_paid_set)
            placeholders = ','.join('?' * len(order_ids_paid_tuple))

            # === ШАГ 2: Получаем детали ТОЛЬКО оплаченных заказов из таблицы `Orders` ===
            query_orders = """
                SELECT OrderID, Name, Email, Phone,DeliveryPoint, ProductName, Quantity, Price, TotalWeight
                FROM Orders
                WHERE LoadedID = ? AND OrderID IN ({})
                ORDER BY OrderID, ProductName -- Сортируем по имени товара для консистентности
                """.format(placeholders)
            # Параметры: loaded_id + список order_id_int
            params_orders = [loaded_id] + list(order_ids_paid_tuple)
            df_orders_paid = pd.read_sql(query_orders, self.conn, params=params_orders)

            if df_orders_paid.empty:
                # Теоретически возможно, если данные в Orders были удалены, но bills остался
                messagebox.showinfo("ℹ️", "Нет данных по оплаченным заказам для печати (детали заказов не найдены).")
                logging.warning(f"Нет данных по оплаченным заказам для загрузки {loaded_id} (детали не найдены).")
                return

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # Заменяем точки в loaded_id на подчеркивания для корректного имени файла
            safe_loaded_id = loaded_id.replace('.', '_')
            docx_filename = f"Заказы_{safe_loaded_id}_{timestamp}.docx"

            # === ШАГ 3: Генерация DOCX ===
            doc = Document()
            # === Установка альбомной ориентации и полей ===
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            # A4 в альбомной ориентации: 297 мм x 210 мм (~11.69 x 8.27 дюймов)
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)
            section.header_distance = Cm(0.5)
            section.footer_distance = Cm(0.5)

            # --- Шрифт по умолчанию ---
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # --- Нижний колонтитул (добавляется один раз, отображается на всех страницах) ---
            footer = section.footer
            p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p_footer.text = f"Файл загрузки: {loaded_name} | Дата загрузки: {loaded_id} | Дата печати: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_footer = p_footer.runs[0]
            run_footer.font.size = Pt(10)

            # --- Для каждого оплаченного заказа ---
            # ... (цикл for начинается) ...
            for order_id, group in df_orders_paid.groupby('OrderID'):
                order_id_int = int(order_id)
                
                # --- ИЗМЕНЕНО: Извлечение данных клиента, включая DeliveryPoint ---
                # Данные клиента (берем из первой строки группы, так как они одинаковые для одного OrderID)
                name = clean_text_for_pdf(group.iloc[0]['Name']) if pd.notna(group.iloc[0]['Name']) else "Не указано"
                email = clean_text_for_pdf(group.iloc[0]['Email']) if pd.notna(group.iloc[0]['Email']) else "Не указано"
                phone = clean_text_for_pdf(group.iloc[0]['Phone']) if pd.notna(group.iloc[0]['Phone']) else "Не указано"
                # --- НОВОЕ: Извлечение точки доставки ---
                delivery_point = clean_text_for_pdf(group.iloc[0]['DeliveryPoint']) if pd.notna(group.iloc[0]['DeliveryPoint']) else "Не указана"
                # --- КОНЕЦ: Извлечение точки доставки ---

                # --- НОВОЕ: Расчет общего веса заказа ---
                # Суммируем TotalWeight по всем строкам одного заказа
                total_order_weight = group['TotalWeight'].sum() if 'TotalWeight' in group.columns else 0.0
                # Убедимся, что это число и обработаем возможные NaN
                if pd.isna(total_order_weight):
                    total_order_weight = 0.0
                # --- КОНЕЦ: Расчет общего веса заказа ---

                # --- НОВОЕ: Расчет количества листов ---
                num_sheets = 1
                if self.config['OPTIONS'].get('MAXWeight', 0) > 0 and total_order_weight > 0:
                    # Импортируем модуль math, если еще не импортирован вверху файла
                    import math
                    # Рассчитываем количество листов, округляя вверх
                    max_weight = self.config['OPTIONS']['MAXWeight']
                    num_sheets = math.ceil(total_order_weight / max_weight)
                # --- КОНЕЦ: Расчет количества листов ---

                # --- НОВОЕ: Цикл по количеству листов ---
                for sheet_number in range(1, num_sheets + 1):
                    
                    # === Таблица 2x1: QR слева, данные справа (повторяется для каждого листа) ===
                    #table_header = doc.add_table(rows=1, cols=2)
                    #table_header.autofit = False
                    # --- Установка ширины столбцов для таблицы заголовка ---
                    # Ширина страницы с учетом полей (~10.5 дюймов)
                    #available_width_header = 14.0
                    # Примерное распределение: QR (2 дюйма), Данные (оставшееся)
                    #table_header.columns[0].width = Cm(5.0) # Ширина QR-кода
                    #table_header.columns[1].width = Cm(available_width_header - 5.0) # Ширина данных клиента

                    #row = table_header.rows[0]
                    
                    # === Таблица 3x1: QR посередине, Лист справа сверху, данные справа снизу ===
                    # Создаем таблицу с 3 строками и 2 столбцами
                    table_header = doc.add_table(rows=2, cols=2)
                    table_header.autofit = False
                    #table_header.style = 'Table Grid' # Уберите, если не нужны рамки
                    # --- Установка ширины столбцов для таблицы заголовка ---
                    available_width_header = 19.0 # или ваше значение в см
                    table_header.columns[0].width = Cm(5.0) # Ширина QR-кода
                    table_header.columns[1].width = Cm(available_width_header - 5.0) # Ширина данных/листа

                    # --- Строка 1: Пустая ячейка слева, Номер листа справа сверху ---
                    cell_sheet_info = table_header.rows[0].cells[1]
                    cell_sheet_info.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    p_sheet_info = cell_sheet_info.paragraphs[0]
                    run_sheet_info = p_sheet_info.add_run(f"Лист {sheet_number} из {num_sheets}")
                    # --- УВЕЛИЧЕНИЕ ШРИФТА ---
                    run_sheet_info.font.size = Pt(24) # Было 12, стало 24. Или Pt(20) для 16.
                    run_sheet_info.bold = True
                    p_sheet_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Выравнивание по правому краю

                    # --- Строка 2: QR-код посередине (объединяем ячейки слева и справа, или просто используем левую) ---
                    # Вариант 1: QR в левой ячейке, центрированный
                    cell_qr = table_header.rows[1].cells[0]
                    cell_qr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell_qr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Центрируем QR в ячейке
                    # Генерируем QR-код (код остается прежним)
                    qr = qrcode.QRCode(version=1, box_size=4, border=4)
                    qr.add_data(order_id_int)
                    qr.make(fit=True)
                    img = qr.make_image(fill_color="black", back_color="white")
                    qr_path = f"temp_qr_order_{order_id_int}.png"
                    img.save(qr_path)
                    p_qr = cell_qr.paragraphs[0]
                    run_qr = p_qr.add_run()
                    # Размер QR-кода чуть меньше высоты строки
                    # Возможно, нужно будет скорректировать размер
                    run_qr.add_picture(qr_path, width=Cm(4.0), height=Cm(4.0)) # Пример размера

                    # --- Строка 3: Пустая ячейка слева, Данные клиента справа снизу ---
                    cell_text = table_header.rows[1].cells[1]
                    cell_text.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    p_text = cell_text.paragraphs[0]
                    p_text.add_run(f"ФИО: {name}").bold = True
                    p_text.add_run(f"\nEmail: {email}")
                    p_text.add_run(f"\nТелефон: {phone}")
                    p_text.add_run(f"\nНомер заказа: {order_id_int}").bold = True
                    # --- Добавление точки доставки и общего веса ---
                    p_text.add_run(f"\nТочка доставки: {delivery_point}")
                    p_text.add_run(f"\nОбщий вес заказа: {total_order_weight:.0f} г")
                    p_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Высота строк (опционально, для лучшего контроля)
                    # table_header.rows[0].height = Cm(1.0) # Высота строки с номером листа
                    # table_header.rows[1].height = Cm(5.0) # Высота строки с QR
                    # table_header.rows[2].height = Cm(3.0) # Высота строки с данными



                    # - Таблица товаров (обновленная) -
                    doc.add_paragraph() # Отступ
                    total_amount = 0.0

                    # --- ИЗМЕНЕНО: Добавлен столбец "Вес" ---
                    # Данные таблицы
                    data = [["Товар", "Кол-во", "Цена", "Стоимость", "Вес (г)"]] # <-- Добавлен столбец
                    for _, row_data in group.iterrows():
                        product = clean_text_for_pdf(row_data['ProductName']) if pd.notna(row_data['ProductName']) else "Не указан"
                        quantity = int(row_data['Quantity']) if pd.notna(row_data['Quantity']) else 0
                        price = float(row_data['Price']) if pd.notna(row_data['Price']) else 0.0
                        line_total = price * quantity
                        total_amount += line_total
                        # --- НОВОЕ: Получение веса позиции ---
                        item_weight = float(row_data['TotalWeight']) if pd.notna(row_data['TotalWeight']) else 0.0
                        # --- КОНЕЦ: Получение веса позиции ---
                        # Добавление строки в таблицу с весом
                        data.append([product, str(quantity), f"{price:.2f}", f"{line_total:.2f}", f"{item_weight:.0f}"]) # <-- Добавлен вес
                    data.append(["ИТОГО ПО ЗАКАЗУ:", "", "", f"{total_amount:.2f} руб.", ""]) # <-- Пустая ячейка для веса в итоговой строке
                    # --- КОНЕЦ: Изменение данных таблицы ---

                    # --- Создание и заполнение таблицы товаров ---
                    # Ширина страницы (в дюймах, с учетом полей ~10.5 дюймов)
                    available_width_items = 14.0
                    # --- ИЗМЕНЕНО: Распределение ширины на 5 столбцов ---
                    # Предположим, распределяем так: Товар(35%), Кол-во(10%), Цена(15%), Стоимость(20%), Вес(20%)
                    column_widths = [
                        8.0,  # Товар
                        1.0,  # Кол-во
                        1.5,  # Цена
                        1.0,  # Стоимость
                        1.5   # Вес (г)
                    ]
                    # Создание таблицы
                    table_items = doc.add_table(rows=1, cols=len(column_widths))
                    table_items.style = 'Table Grid' # Добавляем рамки
                    table_items.autofit = True
                    # Установка ширины столбцов
                    for i, w in enumerate(column_widths):
                        table_items.columns[i].width = Cm(w)

                    # Заполнение таблицы данными
                    hdr_cells = table_items.rows[0].cells
                    # Заголовок
                    for i, text in enumerate(data[0]):
                        hdr_cells[i].text = text
                        hdr_cells[i].paragraphs[0].runs[0].bold = True
                        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # Строки данных
                    for row_data in data[1:]: # Пропускаем заголовок
                        row_cells = table_items.add_row().cells
                        for i, cell_text in enumerate(row_data):
                            row_cells[i].text = cell_text
                            # Выравнивание
                            if i == 0: # Товар - по левому краю
                                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif i == len(row_data) - 1: # Последний столбец (Вес или ИТОГО) - по правому краю
                                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else: # Остальные - по центру
                                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # Форматирование строки "ИТОГО"
                    last_row = table_items.rows[-1].cells
                    for cell in last_row:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    # Выравнивание итоговой суммы по правому краю
                    last_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    last_row[3].paragraphs[0].runs[0].bold = True
                    # Вес в строке итого пустой, выравнивание не критично
                    # --- Конец: Создание и заполнение таблицы товаров ---

                    # Новая страница, если это не последний лист для этого заказа
                    if sheet_number < num_sheets:
                        doc.add_page_break()
                    # Если это последний лист для заказа, но не последний заказ, новая страница добавится в конце внешнего цикла
                # --- КОНЕЦ: Цикл по количеству листов ---
                
                # Удаление временного файла QR-кода для этого заказа (можно сделать один раз в конце, или для каждого листа)
                # Для простоты, удалим в конце внешнего цикла по заказам
                try:
                    # Проверяем, была ли переменная qr_path определена 
                    #(на случай, если group был пуст или возникла ошибка ранее)
                    if 'qr_path' in locals(): 
                        os.remove(qr_path)
                        logging.debug(f"Удалён временный файл QR-кода: {qr_path}")
                    # else: # Не критично, если файл не был создан для этого заказа
                except FileNotFoundError:
                    # Файл уже был удален или не существовал
                    logging.debug(f"Временный файл QR-кода {qr_path if 'qr_path' in locals() else 'неизвестно'} не найден при удалении.")
                except Exception as e:
                # Другая ошибка при удалении
                    logging.warning(f"Не удалось удалить временный файл QR-кода {qr_path if 'qr_path' in locals() else 'неизвестно'}: {e}")
 
                
            # ... (цикл for заканчивается) ...




            # === ШАГ 4: Добавление таблицы НЕОПЛАЧЕННЫХ/ОТМЕНЕННЫХ заказов ===
            # Выбираем счета из bills с тем же loaded_id, но со статусом НЕ 'succeeded'
            query_unpaid_bills = """
                SELECT order_id_int, name, email, total_amount, status
                FROM bills
                WHERE loaded_id = ? AND status <> 'succeeded' -- Или конкретные статусы IN ('pending', 'cancelled')
            """
            df_unpaid_bills = pd.read_sql(query_unpaid_bills, self.conn, params=[loaded_id])

            if not df_unpaid_bills.empty:
                doc.add_page_break() # Новая страница

                # Добавляем заголовок
                p = doc.add_paragraph()
                p.add_run("Неоплаченные и отмененные заказы").bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Создаем таблицу
                table_uc = doc.add_table(rows=1, cols=5)
                table_uc.style = 'Table Grid'
                hdr_cells_uc = table_uc.rows[0].cells
                hdr_texts_uc = ["Номер заказа", "ФИО", "Email", "Сумма", "Статус"]
                for i, text in enumerate(hdr_texts_uc):
                    hdr_cells_uc[i].text = text
                    hdr_cells_uc[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    hdr_cells_uc[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    hdr_cells_uc[i].paragraphs[0].runs[0].bold = True

                # Заполняем таблицу данными
                for _, row in df_unpaid_bills.iterrows():
                    row_cells = table_uc.add_row().cells
                    # Обязательно обрабатываем возможные None
                    order_id_uc = int(row['order_id_int']) if pd.notna(row['order_id_int']) else "N/A"
                    name_uc = clean_text_for_pdf(row['name']) if pd.notna(row['name']) else "Не указано"
                    email_uc = clean_text_for_pdf(row['email']) if pd.notna(row['email']) else "Не указано"
                    amount_uc = f"{float(row['total_amount']):.2f} руб." if pd.notna(row['total_amount']) else "N/A"
                    status_uc = row['status'] if pd.notna(row['status']) else "N/A"

                    row_cells[0].text = str(order_id_uc)
                    row_cells[1].text = name_uc
                    row_cells[2].text = email_uc
                    row_cells[3].text = amount_uc
                    row_cells[4].text = status_uc

                    # Форматирование ячеек (выравнивание и т.д.)
                    for i in range(5):
                        row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        if i < 3: # Номер, ФИО, Email - по левому краю
                            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else: # Сумма, Статус - по правому краю
                            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # === Сохранение DOCX ===
            doc.save(docx_filename)
            logging.info(f"DOCX сохранён: {docx_filename}")

            # Удаление временных QR-файлов
            #for file in os.listdir('.'):
            #    if file.startswith("temp_qr_order_") and file.endswith(".png"):
            #        try:
            #            os.remove(file)
            #            logging.debug(f"Удалён временный файл QR-кода: {file}")
            #        except Exception as e:
            #            logging.warning(f"Не удалось удалить временный файл {file}: {e}")

            messagebox.showinfo("✅", f"DOCX успешно сохранён:\n{docx_filename}")

        except Exception as e:
            error_msg = f"Ошибка генерации DOCX: {e}"
            logging.error(error_msg, exc_info=True) # exc_info=True для полного трейса
            messagebox.showerror("❌ Ошибка", error_msg)

    # ==================== ФУНКЦИИ ДЛЯ СЧЕТОВ ====================
    
    def create_bills_widgets(self):
        """Создание элементов для вкладки счетов"""
        # Основной контейнер с прокруткой
        canvas = tk.Canvas(self.bills_frame)
        scrollbar = ttk.Scrollbar(self.bills_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # --- Секция статуса API ---
        api_frame = ttk.LabelFrame(scrollable_frame, text="Статус API ЮKassa", padding=10)
        api_frame.pack(fill=tk.X, pady=(0, 10), padx=5)
        self.api_status_label = ttk.Label(api_frame, text="Проверка...", foreground="orange")
        self.api_status_label.pack(side=tk.LEFT)

        # --- Секция обновления статусов ---
        update_frame = ttk.LabelFrame(scrollable_frame, text="Обновление статусов счетов", padding=10)
        update_frame.pack(fill=tk.X, pady=(0, 10), padx=5)
        self.btn_get_info = ttk.Button(
            update_frame,
            text="🔄 Обновить статусы счетов",
            command=self.get_bills_info,
            state=tk.DISABLED
        )
        self.btn_get_info.pack(pady=5)

        # --- Область для вывода результатов статистики ---
        # Создаем фрейм для области результатов, НЕ растягиваем его по вертикали
        self.result_frame_stats = tk.Frame(update_frame)
        # Используем fill=tk.X, чтобы растягивался по ширине, но expand=False (или его отсутствие)
        # чтобы не занимал всё доступное вертикальное пространство
        self.result_frame_stats.pack(pady=10, padx=10, fill=tk.X) # Изменено: убран expand=True

        # Создаем Text с фиксированной высотой в 7 строк
        self.result_text_stats = tk.Text(
            self.result_frame_stats,
            wrap=tk.WORD,
            state=tk.DISABLED,
            height=7  # <-- Установка высоты в 7 строк
        )
        # Создаем Scrollbar для Text
        self.result_scrollbar_stats = tk.Scrollbar(
            self.result_frame_stats,
            orient="vertical",
            command=self.result_text_stats.yview
        )
        # Связываем Scrollbar с Text
        self.result_text_stats.configure(yscrollcommand=self.result_scrollbar_stats.set)
        # Упаковываем Text и Scrollbar внутри их фрейма
        # Text растягивается по ширине и заполняет фрейм по вертикали (высота фиксирована самим Text)
        self.result_text_stats.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        # Scrollbar растягивается по вертикали (в пределах высоты Text)
        self.result_scrollbar_stats.pack(side=tk.RIGHT, fill=tk.Y)
        # --- Конец области для вывода результатов ---

        # --- Секция автоматического создания счетов ---
        auto_create_frame = ttk.LabelFrame(scrollable_frame, text="Автоматическое создание счетов", padding=10)
        auto_create_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10), padx=5)
        ttk.Label(auto_create_frame, text="Эта функция создаст счета в ЮKassa для всех заказов,\nу которых в базе данных bill_id отсутствует.").pack(pady=(0, 10))
        self.btn_auto_create = ttk.Button(
            auto_create_frame,
            text="Создать счета для новых заказов",
            command=self.auto_create_invoices,
            state=tk.DISABLED
        )
        self.btn_auto_create.pack(pady=5)
        # --- Прогресс и логи ---
        self.progress_frame = ttk.Frame(auto_create_frame)
        # progress_bar и progress_label будут созданы позже при необходимости
        self.log_frame = ttk.Frame(auto_create_frame)
        self.log_text_widget = scrolledtext.ScrolledText(self.log_frame, wrap=tk.WORD, state=tk.DISABLED, height=15)
        self.log_text_widget.pack(fill=tk.BOTH, expand=True)
 
 
 
 
    def log_to_widget(self, message):
        """Логирование операции создания счетов в виджет GUI"""
        self.log_text_widget.config(state=tk.NORMAL)
        self.log_text_widget.insert(tk.END, f"[{datetime.now().strftime('%H:%M:%S')}] {message}\n")
        self.log_text_widget.see(tk.END)
        self.log_text_widget.config(state=tk.DISABLED)
        self.root.update_idletasks()

    def update_invoice_statuses(self):
        """Обновление статусов счетов через API ЮKassa"""
        if not self.conn:
            messagebox.showwarning("⚠️", "Нет подключения к БД.")
            return False
        yookassa_config = self.config.get('YOOCASSA', {})
        shop_id = yookassa_config.get('shop_id', '')
        secret_key = yookassa_config.get('secret_key', '')
        if not shop_id or not secret_key:
            messagebox.showwarning("⚠️", "API ЮKassa не настроен.")
            return False
        try:
            cursor = self.conn.cursor()
            # Выбираем только те счета, у которых есть bill_id (invoice_id)
            cursor.execute("SELECT id, bill_id FROM bills WHERE bill_id IS NOT NULL AND bill_id <> ''")
            rows = cursor.fetchall()
            if not rows:
                logging.info("Нет счетов с bill_id для обновления статусов.")
                return True # Нечего обновлять, но это не ошибка
            updated_count = 0
            error_count = 0
            for bill_id_db, invoice_id in rows:
                try:
                    # 1. Формируем URL для получения информации о счете
                    url = f"https://api.yookassa.ru/v3/invoices/{invoice_id}"
                    # 2. Подготавливаем заголовки аутентификации
                    credentials = f"{shop_id}:{secret_key}"
                    encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')
                    headers = {
                        'Authorization': f"Basic {encoded_credentials}",
                        'Content-Type': 'application/json'
                    }
                    # 3. Отправляем GET-запрос
                    logging.debug(f"Отправка GET-запроса на {url}")
                    response = requests.get(url, headers=headers)
                    # 4. Обрабатываем ответ
                    if response.status_code == 200:
                        invoice_data = response.json()
                        new_status = invoice_data.get('status')
                        if new_status:
                            # 5. Обновляем статус в БД
                            update_cursor = self.conn.cursor()
                            # Используем параметризованный запрос для безопасности
                            update_cursor.execute(
                                "UPDATE bills SET status = ? WHERE id = ?",
                                new_status, bill_id_db
                            )
                            if update_cursor.rowcount > 0:
                                self.conn.commit()
                                updated_count += 1
                                logging.info(f"Статус счета {invoice_id} обновлен до '{new_status}' (DB ID: {bill_id_db})")
                            else:
                                logging.warning(f"Статус счета {invoice_id} не был обновлен (DB ID: {bill_id_db}), возможно, статус не изменился или запись не найдена.")
                            update_cursor.close()
                        else:
                           logging.warning(f"Не удалось получить статус из ответа для счета {invoice_id} (DB ID: {bill_id_db}). Ответ: {invoice_data}")
                           error_count += 1
                    elif response.status_code == 404:
                        # Счет не найден - возможно, он был удален или ID неверен
                        logging.warning(f"Счет {invoice_id} не найден в ЮKassa (404) (DB ID: {bill_id_db}).")
                        # Можно решить, обновлять ли статус на какой-то специальный, например 'not_found'
                        # update_cursor = self.conn.cursor()
                        # update_cursor.execute("UPDATE bills SET status = ? WHERE id = ?", 'not_found', bill_id_db)
                        # self.conn.commit()
                        # update_cursor.close()
                        error_count += 1
                    else:
                        # Другая ошибка API
                        try:
                            error_data = response.json()
                            error_msg_detail = json.dumps(error_data, ensure_ascii=False)
                        except:
                            error_msg_detail = response.text
                        logging.error(f"Ошибка API при получении статуса счета {invoice_id} (DB ID: {bill_id_db}) (HTTP {response.status_code}): {error_msg_detail}")
                        error_count += 1
                except requests.exceptions.RequestException as e:
                    # Ошибки сети, таймауты и т.д.
                    logging.error(f"Ошибка сети при обновлении статуса счета {invoice_id} (DB ID: {bill_id_db}): {e}")
                    error_count += 1
                except Exception as e:
                    # Другие исключения
                    logging.error(f"Исключение при обновлении статуса счета {invoice_id} (DB ID: {bill_id_db}): {e}", exc_info=True)
                    error_count += 1
            logging.info(f"Обновление статусов завершено. Успешно: {updated_count}, Ошибок: {error_count}")
            if error_count > 0:
                 messagebox.showwarning("⚠️ Обновление статусов", f"Во время обновления статусов произошли ошибки.\nУспешно обновлено: {updated_count}\nОшибок: {error_count}")
            else:
                 messagebox.showinfo("✅ Обновление статусов", f"Статусы счетов успешно обновлены.\nОбновлено записей: {updated_count}")
            return True
        except Exception as e:
            error_msg = f"Критическая ошибка при обновлении статусов счетов: {e}"
            logging.error(error_msg, exc_info=True)
            messagebox.showerror("❌ Ошибка", error_msg)
            return False

    def get_bills_info(self):
        """Получение и отображение статистики счетов (с обновлением статусов)"""
        if not self.conn:
            messagebox.showwarning("⚠️", "Нет подключения к БД.")
            return
        # --- НОВОЕ: Обновляем статусы перед отображением статистики ---
        self.log_to_widget("🔄 Начинаем обновление статусов счетов...")
        update_success = self.update_invoice_statuses()
        if not update_success:
            # Если обновление не удалось, всё равно покажем старую статистику
            self.log_to_widget("❌ Обновление статусов не удалось, показываем текущую статистику.")
        else:
             self.log_to_widget("✅ Обновление статусов завершено.")
        try:
            cursor = self.conn.cursor()
            # 1. Счета без bill_id
            cursor.execute("SELECT COUNT(*) FROM bills WHERE bill_id IS NULL OR bill_id = ''")
            count_no_bill_id = cursor.fetchone()[0]
            # 2. Счета со статусом 'pending'
            cursor.execute("SELECT COUNT(*) FROM bills WHERE status = ?", 'pending')
            count_pending = cursor.fetchone()[0]
            # 3. Счета со статусом 'cancelled'/'canceled'
            cursor.execute("SELECT COUNT(*) FROM bills WHERE status IN (?, ?)", ('cancelled', 'canceled'))
            count_cancelled = cursor.fetchone()[0]
            # --- НОВОЕ: Добавляем счета со статусом 'paid' ---
            cursor.execute("SELECT COUNT(*) FROM bills WHERE status = ?", 'succeeded')
            count_paid = cursor.fetchone()[0]
            # --- НОВОЕ: Добавляем счета со статусом 'waiting_for_capture' ---
            cursor.execute("SELECT COUNT(*) FROM bills WHERE status = ?", 'waiting_for_capture')
            count_waiting_for_capture = cursor.fetchone()[0]
            # --- Формирование результата ---
            result_text = (
                f"Счетов без bill_id: {count_no_bill_id}\n"
                f"Счетов со статусом 'pending': {count_pending}\n"
                f"Счетов со статусом 'waiting_for_capture': {count_waiting_for_capture}\n" # Новый статус
                f"Счетов со статусом 'succeeded': {count_paid}\n" # Новый статус
                f"Счетов со статусом 'cancelled/canceled': {count_cancelled}\n"
                "---\n"
                )
            # --- Отображение результата ---
            self.result_text_stats.config(state=tk.NORMAL)
            self.result_text_stats.delete(1.0, tk.END)
            self.result_text_stats.insert(tk.END, result_text)
            self.result_text_stats.config(state=tk.DISABLED)
            logging.info("Статистика счетов (с учетом обновленных статусов) получена и отображена.")
        except Exception as e:
            error_msg = f"Ошибка при получении статистики: {e}"
            logging.error(error_msg)
            messagebox.showerror("❌ Ошибка", error_msg)
            self.result_text_stats.config(state=tk.NORMAL)
            self.result_text_stats.delete(1.0, tk.END)
            self.result_text_stats.insert(tk.END, f"Ошибка: {error_msg}\n")
            self.result_text_stats.config(state=tk.DISABLED)

    def auto_create_invoices(self, event_datetime=None):
        """Автоматическое создание счетов для заказов без bill_id"""
        if not self.conn:
            messagebox.showwarning("⚠️", "Нет подключения к БД.")
            return
        yookassa_config = self.config.get('YOOCASSA', {})
        shop_id = yookassa_config.get('shop_id', '')
        secret_key = yookassa_config.get('secret_key', '')
        if not shop_id or not secret_key:
            messagebox.showwarning("⚠️", "API ЮKassa не настроен.")
            return
        try:
            cursor = self.conn.cursor()
            # 1. Получаем заказы без bill_id
            cursor.execute("SELECT id, order_id_int, name, email, total_amount FROM bills WHERE bill_id IS NULL OR bill_id = ''")
            rows = cursor.fetchall()
            if not rows:
                messagebox.showinfo("ℹ️ Информация", "Нет заказов без bill_id для создания счетов.")
                return
            total_orders = len(rows)
            self.log_to_widget(f"Найдено {total_orders} заказов для создания счетов.")
            # --- Настройка GUI для отображения прогресса ---
            if not hasattr(self, 'progress_bar'):
                self.progress_frame.pack(fill=tk.X, pady=(10, 5))
                self.progress_var = tk.DoubleVar()
                self.progress_bar = ttk.Progressbar(self.progress_frame, variable=self.progress_var, maximum=total_orders)
                self.progress_bar.pack(fill=tk.X)
                self.progress_label = ttk.Label(self.progress_frame, text="0 / 0")
                self.progress_label.pack()
            if not hasattr(self, 'log_frame_packed'):
                 self.log_frame.pack(fill=tk.BOTH, expand=True, pady=(5, 0))
                 self.log_text_widget.pack(fill=tk.BOTH, expand=True)
                 self.log_frame_packed = True
            self.log_text_widget.config(state=tk.NORMAL)
            self.log_text_widget.delete(1.0, tk.END)
            self.log_text_widget.config(state=tk.DISABLED)
            self.btn_auto_create.config(state=tk.DISABLED, text="Создание счетов...")
            self.root.update_idletasks()
            # --- Конец настройки GUI ---
            success_count = 0
            error_count = 0
            for i, (bill_id_db, order_id_int, name, email, total_amount) in enumerate(rows):
                try:
                    # Обновление прогресса
                    self.progress_var.set(i + 1)
                    self.progress_label.config(text=f"{i + 1} / {total_orders}")
                    self.root.update_idletasks()
                    # 2. Формируем данные для API
                    price_value = round(float(total_amount), 2)
                    # Описание счета: ФИО клиента и номер заказа
                    description = f"{name}, заказ №{order_id_int}"
                    # --- НОВОЕ: Срок действия - до 23:00 MSK текущей даты ---
                    # 1. Получаем текущую дату и время в Московском часовом поясе (MSK)
                    msk_tz = pytz.timezone('Europe/Moscow')
                    now_msk = datetime.now(msk_tz)

                    # 2. Создаем объект datetime для сегодняшней даты, 23:00 по MSK
                    # replace(hour=23, minute=0, second=0, microsecond=0) устанавливает время
                    # astimezone(msk_tz) гарантирует, что объект знает о часовом поясе MSK
                    expires_at_msk = now_msk.replace(hour=23, minute=0, second=0, microsecond=0).astimezone(msk_tz)

                    # 3. Преобразуем время истечения из MSK в UTC, как требует API ЮKassa
                    utc_tz = pytz.UTC
                    expires_at_utc = expires_at_msk.astimezone(utc_tz)

                    # 4. Форматируем время в UTC в строку в формате, ожидаемом API
                    # Формат: YYYY-MM-DDTHH:MM:SS.fffZ
                    expires_at_str = expires_at_utc.strftime('%Y-%m-%dT%H:%M:%S.%f')[:-3] + 'Z'
                    # --- КОНЕЦ: Новый срок действия ---
                    # Формируем тело запроса для /v3/invoices
                    invoice_data = {
                        "cart": [
                            {
                                "description": "Внесение депозита", # Фиксированное описание товара
                                "price": {
                                    "value": f"{price_value:.2f}",
                                    "currency": "RUB"
                                },
                                "quantity": 1.0 # Фиксированное количество
                            }
                        ],
                        "description": description,
                        "expires_at": expires_at_str,
                        "payment_data": {
                            "amount": {
                                "value": f"{price_value:.2f}",
                                "currency": "RUB"
                            },
                            "capture": True # Автоподтверждение
                        },
                        "delivery_method": {
                            "type": "self" # Самостоятельная доставка (ссылка на оплату)
                        },
                        "metadata": {
                            "order_id": str(order_id_int),
                            "internal_bill_id": str(bill_id_db) # На случай, если понадобится сопоставить
                        }
                    }
                    # Добавляем receipt, если есть email и ФИО
                    if email and '@' in email and name:
                        receipt_items = [{
                            "description": "Внесение депозита", # Фиксированное описание товара
                            "quantity": 1.0, # Фиксированное количество
                            "amount": {
                                "value": f"{price_value:.2f}",
                                "currency": "RUB"
                            },
                            "vat_code": yookassa_config['vat_code'], # Налоговая ставка из конфигурации
                            "payment_subject": "payment", # Предмет расчета
                            "payment_mode": "advance" # Способ расчета
                        }]
                        invoice_data["payment_data"]["receipt"] = {
                            "items": receipt_items,
                            "tax_system_code": yookassa_config['tax_system_code'], # Система налогообложения из конфигурации
                            "customer": {
                                "full_name": name, # ФИО клиента
                                "email": email # Email клиента
                            }
                        }
                    # 3. Отправляем запрос
                    url = "https://api.yookassa.ru/v3/invoices"
                    credentials = f"{shop_id}:{secret_key}"
                    encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')
                    headers = {
                        'Authorization': f"Basic {encoded_credentials}",
                        'Content-Type': 'application/json',
                        'Idempotence-Key': str(uuid.uuid4())
                    }
                    # --- ИСПРАВЛЕНО: ЛОГИРОВАНИЕ ЗАПРОСА ---
                    logging.info(f"Запрос к API: {{'method': 'POST', 'url': '{url}', 'headers': {{...}}, 'data': {json.dumps(invoice_data, ensure_ascii=False)}}}")
                    self.log_to_widget(f"🌐 Отправка запроса на создание счёта для заказа {order_id_int}...")
                    # --- КОНЕЦ: ЛОГИРОВАНИЕ ЗАПРОСА ---
                    response = requests.post(url, headers=headers, json=invoice_data)
                    # --- ИСПРАВЛЕНО: ЛОГИРОВАНИЕ ОТВЕТА ---
                    try:
                        response_content = response.json()
                    except:
                        response_content = response.text
                    response_log_data = {
                        "status_code": response.status_code,
                        "headers": dict(response.headers),
                        "content": response_content
                    }
                    logging.info(f"Ответ от API: {json.dumps(response_log_data, ensure_ascii=False, indent=2)}")
                    # --- КОНЕЦ: ЛОГИРОВАНИЕ ОТВЕТА ---
                    # 4. Обрабатываем ответ
                    if response.status_code in [200, 201]:
                        result = response.json()
                        yookassa_invoice_id = result.get('id')
                                                
                        # --- НОВОЕ: Извлечение и преобразование expires_at ---
                        

                        expires_at_formatted_for_db = None # Для сохранения в БД, если нужно
                        expires_at_formatted_for_email = "не указан" # Для вставки в письмо

                        yookassa_expires_at_raw = result.get('expires_at')
                        if yookassa_expires_at_raw:
                            try:
                                # 1. Парсим строку времени из API (она в UTC)
                                # Формат: 2025-08-14T23:19:28.320Z
                                expires_at_utc = datetime.fromisoformat(yookassa_expires_at_raw.replace('Z', '+00:00'))
                                
                                # 2. Преобразуем из UTC в Московское время (MSK)
                                utc_tz = pytz.utc
                                msk_tz = pytz.timezone('Europe/Moscow')
                                # Убедимся, что объект datetime "знает" о временной зоне UTC
                                expires_at_utc_aware = utc_tz.localize(expires_at_utc) if expires_at_utc.tzinfo is None else expires_at_utc
                                # Преобразуем в MSK
                                expires_at_msk = expires_at_utc_aware.astimezone(msk_tz)
                                
                                # 3. Форматируем для отображения в письме (HH:MM dd.MM.yyyy)
                                expires_at_formatted_for_email = expires_at_msk.strftime('%H:%M %d.%m.%Y')
                                
                                # 4. (Опционально) Форматируем для сохранения в БД (если добавите столбец)
                                # expires_at_formatted_for_db = expires_at_msk # Сохраняем как datetime объект
                                
                            except Exception as e_parse_expires:
                                logging.error(f"Ошибка при парсинге/преобразовании expires_at '{yookassa_expires_at_raw}': {e_parse_expires}")
                                # expires_at_formatted_for_email останется "не указан"
                        # --- КОНЕЦ: Извлечение и преобразование expires_at ---
                        
                        # --- ИСПРАВЛЕНО: ИЗВЛЕЧЕНИЕ И ОБРЕЗКА URL ---
                        payment_url_raw = None
                        payment_url_final = None
                        delivery_method = result.get('delivery_method', {})
                        if isinstance(delivery_method, dict):
                            payment_url_raw = delivery_method.get('url')
                        if payment_url_raw:
                            # Убираем '/a' в конце, если он есть
                            if payment_url_raw.endswith('/a'):
                                payment_url_final = payment_url_raw[:-2]
                            elif payment_url_raw.endswith('/a  '): # Иногда в API попадают лишние пробелы
                                payment_url_final = payment_url_raw[:-4]
                            else:
                                payment_url_final = payment_url_raw.rstrip() # Просто убираем пробелы справа на всякий случай
                        if not payment_url_final:
                            logging.warning(f"URL для оплаты не найден или не может быть обработан для счета {yookassa_invoice_id} (заказ {order_id_int}). Raw URL: '{payment_url_raw}'")
                            self.log_to_widget(f"⚠️ URL оплаты не обработан для счета {yookassa_invoice_id}.")
                        else:
                            self.log_to_widget(f"✅ URL оплаты: {payment_url_final}")
                        # --- КОНЕЦ: ИЗВЛЕЧЕНИЕ И ОБРЕЗКА URL ---
                        if yookassa_invoice_id:
                            # 5. Обновляем запись в БД (включая URL)
                            update_cursor = self.conn.cursor()
                            # --- ИСПРАВЛЕНО: Обновляем и bill_id, и urls ---
                            update_cursor.execute(
                                "UPDATE bills SET bill_id = ?, urls = ?, expires_at = ? WHERE id = ?",
                                yookassa_invoice_id, payment_url_final, expires_at_msk, bill_id_db # Используем первичный ключ `id` таблицы bills
                            )
                            # --- КОНЕЦ: Обновление bill_id и urls ---
                            self.conn.commit()
                            update_cursor.close()
                            self.log_to_widget(f"✅ Счет создан для заказа {order_id_int} (ID: {yookassa_invoice_id})")
                            # --- НОВОЕ: Отправка email ---
                            # Получаем обновленный URL из БД
                            fetch_url_cursor = self.conn.cursor()
                            fetch_url_cursor.execute("SELECT urls FROM bills WHERE id = ?", bill_id_db)
                            row_url = fetch_url_cursor.fetchone()
                            payment_url_for_email = row_url[0] if row_url else None
                            fetch_url_cursor.close()
                            # Отправляем email
                            email_sent = self.send_payment_email(
                                order_id_int=int(order_id_int),
                                client_name=name,
                                client_email=email,
                                payment_url_db=payment_url_for_email,
                                expires_at_str=expires_at_formatted_for_email
                            )
                            if email_sent:
                                self.log_to_widget(f"📧 Письмо отправлено клиенту {name} ({email})")
                            else:
                                self.log_to_widget(f"❌ Ошибка отправки email клиенту {name} ({email})")
                            # --- КОНЕЦ: Отправка email ---
                            success_count += 1
                        else:
                            self.log_to_widget(f"❌ Ошибка: Не получен ID счета от ЮKassa для заказа {order_id_int}. Ответ: {result}")
                            error_count += 1
                    else:
                        # Обрабатываем ошибку API
                        try:
                            error_data = response.json()
                            error_msg_detail = json.dumps(error_data, ensure_ascii=False)
                        except:
                            error_msg_detail = response.text
                        self.log_to_widget(f"❌ Ошибка API для заказа {order_id_int} (HTTP {response.status_code}): {error_msg_detail[:200]}...")
                        logging.error(f"Ошибка API для заказа {order_id_int} (HTTP {response.status_code}): {error_msg_detail}")
                        error_count += 1
                except Exception as e:
                    self.log_to_widget(f"❌ Исключение при обработке заказа {order_id_int}: {e}")
                    logging.error(f"Исключение при обработке заказа {order_id_int}: {e}", exc_info=True)
                    error_count += 1
                    # Продолжаем обработку следующих заказов
            # Завершение прогресса
            self.progress_var.set(total_orders)
            self.progress_label.config(text=f"{total_orders} / {total_orders}")
            # Финальное сообщение
            final_message = f"Создание счетов завершено.\nУспешно: {success_count}\nОшибок: {error_count}"
            self.log_to_widget(final_message)
            messagebox.showinfo("🏁 Завершено", final_message)
        except Exception as e:
            error_msg = f"Критическая ошибка в процессе автоматического создания счетов: {e}"
            logging.error(error_msg, exc_info=True)
            self.log_to_widget(f"❌ {error_msg}")
            messagebox.showerror("❌ Критическая ошибка", error_msg)
        finally:
            self.btn_auto_create.config(state=tk.NORMAL, text="Создать счета для новых заказов")

    def send_payment_email(self, order_id_int, client_name, client_email, payment_url_db, expires_at_str):
        """
        Отправляет email с ссылкой на оплату и QR-кодом через Exchange 2019.
        :param order_id_int: Номер заказа (int)
        :param client_name: ФИО клиента (str)
        :param client_email: Email клиента (str)
        :param payment_url_db: URL из базы данных (str или None)
        :param expires_at_str: Отформатированная дата истечения (str)
        :return: True если успешно, False если ошибка
        """
        if not self.conn:
            logging.error("Нет подключения к БД для отправки email.")
            return False
        try:
            # 1. Проверяем и формируем URL-ы
            if not payment_url_db:
                logging.error(f"URL оплаты для заказа {order_id_int} отсутствует в БД.")
                return False
            # URL для перехода по ссылке (добавляем /l)
            payment_url_link = payment_url_db.rstrip('/') + "/l"
            # URL для генерации QR-кода (добавляем /q)
            payment_url_qr = payment_url_db.rstrip('/') + "/q"
            # 2. Загружаем конфигурацию почты
            mail_config = self.config.get('MAIL', {})
            # Проверяем обязательные поля
            required_fields = ['smtp_server']
            missing_fields = [field for field in required_fields if not mail_config.get(field)]
            if missing_fields:
                logging.error(f"В конфигурации почты отсутствуют обязательные поля: {missing_fields}")
                return False
            # 3. Формируем HTML-тело письма на основе шаблона
            # Предполагаем, что total_amount передаётся в функцию или получается из БД
            # Для получения суммы, сделаем запрос к БД
            try:
                amount_cursor = self.conn.cursor()
                amount_cursor.execute("SELECT total_amount FROM bills WHERE order_id_int = ? AND email = ?", order_id_int, client_email)
                amount_row = amount_cursor.fetchone()
                order_amount = float(amount_row[0]) if amount_row and amount_row[0] is not None else 0.0
                amount_cursor.close()
            except Exception as e:
                logging.warning(f"Не удалось получить сумму заказа {order_id_int} для email {client_email}: {e}. Используется 0.0.")
                order_amount = 0.0
            # --- ОБНОВЛЕНИЕ: Используем CID для изображений ---
            html_body = f"""
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Strict//EN" "http://www.w3.org/TR/xhtml1/DTD/xhtml1-strict.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
  <meta http-equiv="Content-Type" content="text/html; charset=utf-8" />
  <meta name="viewport" content="width=device-width"/>
  <title>Счет на оплату</title>
</head>
<body style="width: 100% !important; min-width: 100%; -webkit-text-size-adjust: 100%; -ms-text-size-adjust: 100%; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; text-align: left; line-height: 1.3; font-size: 16px; margin: 0; padding: 0;" align="left">
<table class="body" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; height: 100%; width: 100%; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; background: #f0f0f0; padding: 0;" width="100%" bgcolor="#f0f0f0">
  <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
    <td class="center" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: center; margin: 0; padding: 0;" align="center" valign="top">
      <center style="width: 100%; min-width: 580px;">
      <table class="container" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: inherit; width: 580px; background: #fefefe; margin: 0 auto; padding: 0;">
        <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
          <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
            <!-- === НАЧАЛО: ШАПКА С ЛОГОТИПОМ === -->
                <table class="row" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 100%; display: block; padding: 0;">
                  <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                    <td class="wrapper last" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; position: relative; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 10px 0 0;" align="left" valign="top">
                      <table class="twelve columns" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 580px; margin: 0 auto; padding: 0;">
                        <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                          <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0 20px 10px;" align="left" valign="top">
                            <!-- Логотип и информация об организации -->
                            <table style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 100%; padding: 50px 0 20px;">
                              <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                                <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0 0 10px;" align="center" valign="top">
                                  <!-- Изображение с шириной 540px (580px контейнер - 20px*2 padding) и автоматической высотой -->
                                  <img src="cid:logo1"
                                       style="display: block; margin: 0 auto; width: 540px; height: auto; max-width: 100%;"
                                       width="540" /><br>
                                </td>
                              </tr>
                            </table>
                          </td>
                        </tr>
                      </table>
                    </td>
                  </tr>
                </table>
                <!-- === КОНЕЦ: ШАПКА С ЛОГОТИПОМ === -->
            <table class="row" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 100%; display: block; padding: 0;">
              <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                <td class="wrapper last" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; position: relative; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                  <table class="twelve columns" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 580px; margin: 0 auto; padding: 0;">
                    <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                      <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0 20px 10px;" align="left" valign="top">
                        <p style="color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; text-align: left; line-height: 1.3; margin: 0 0 0; padding: 30px 0 0;" align="left"><span style="font-size: 24px;"><strong>Добрый день, </strong></span></p>
                        <p style="color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; text-align: left; line-height: 1.3; margin: 0 0 10px; padding: 0;" align="left"><span style="font-size: 24px;"><strong>{client_name}</strong></span></p>
                      </td>
                    </tr>
                  </table>
                </td>
              </tr>
            </table>
            <table class="row" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 100%; display: block; padding: 0;">
              <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                <td class="wrapper last" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; position: relative; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                  <table class="twelve columns" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 580px; margin: 0 auto; padding: 0;">
                    <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                      <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0 20px 10px;" align="left" valign="top">
                        <p style="text-align: left; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; line-height: 1.3; margin: 0 0 10px; padding: 20px 0 0;" align="left">Спасибо, что оформили заказ на нашем сайте!</p>
                      </td>
                    </tr>
                  </table>
                </td>
              </tr>
            </table>
            <!-- === НОВОЕ: НОМЕР ЗАКАЗА И СУММА В ВИДЕ ТАБЛИЦЫ 2x2 === -->
            <table class="row" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 100%; display: block; padding: 0; margin-top: 10px;">
              <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                <td class="wrapper" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; position: relative; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                  <table class="twelve columns" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 580px; margin: 0 auto; padding: 0;">
                    <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                      <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0 20px 10px;" align="left" valign="top">
                        <!-- Таблица 2x2 для номера заказа и суммы -->
                        <table style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 45%; margin: 0; padding: 0;">
                          <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                            <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                              <strong>Заказ №</strong>
                            </td>
                            <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                              {int(order_id_int)}
                            </td>
                          </tr>
                          <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                            <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                              <strong>Сумма</strong>
                            </td>
                            <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                              {order_amount:.2f} руб.
                            </td>
                          </tr>
                        </table>
                      </td>
                    </tr>
                  </table>
                </td>
              </tr>
            </table>
            <!-- === КОНЕЦ: НОМЕР ЗАКАЗА И СУММА В ВИДЕ ТАБЛИЦЫ 2x2 === -->
            <table class="row" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 100%; display: block; padding: 0; margin-top: 10px;">
              <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                <td class="wrapper last" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; position: relative; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                  <table class="twelve columns" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 580px; margin: 0 auto; padding: 0;">
                    <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                      <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0 20px 10px;" align="left" valign="top">
                        <p style="text-align: left; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; line-height: 1.3; margin: 0 0 10px; padding: 0;" align="left">
                          Чтобы завершить процесс заказа, пожалуйста, произведите оплату по ссылке ниже до 
                          <span style="font-weight: bold;">{expires_at_str}</span> г.
                          <br>
                          После указанного срока ссылка будет неактивна.
                        </p> 
                      </td>
                    </tr>
                  </table>
                </td>
              </tr>
            </table>
            <table class="row" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 100%; display: block; padding: 0;">
              <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                <td class="wrapper last" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; position: relative; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                  <table class="twelve columns" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 580px; margin: 0 auto; padding: 0;">
                    <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                      <td class="center" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: center; margin: 0; padding: 0 20px 10px;" align="center" valign="top">
                        <!--[if mso]>
                        <p style='line-height:0;margin:0;'>&nbsp;</p>
                        <v:roundrect arcsize='25%' fill='t' fillcolor='#f5b333' href='{payment_url_link}' stroke='f' strokecolor='' strokeweight='1px' style='v-text-anchor:middle;width:500px;height:45px;mso-padding-alt:0;padding:13px 20px;' xmlns:v='urn:schemas-microsoft-com:vml' xmlns:w='urn:schemas-microsoft-com:office:word'>
                          <w:anchorlock />
                          <center style='color: #FFF; font-family:sans-serif; font-size:16px; font-weight:bold; mso-line-height-rule:exactly; mso-text-raise:4px'>Оплатить заказ</center>
                        </v:roundrect>
                        <![endif]-->
                        <!--[if !mso]><!-- -->
                        <a href="{payment_url_link}" style="line-height: 16px; font-size: 16px !important; display: block; width: auto; border-radius: 25px; -webkit-border-radius: 25px; -moz-border-radius: 25px; color: #ffffff; text-decoration: none; font-weight: bold; font-family: Arial, sans-serif; text-align: center; height: 100%; background: #f5b333 repeat center center; padding: 13px 20px;">Оплатить заказ</a>
                        <!--<![endif]-->
                      </td>
                    </tr>
                  </table>
                </td>
              </tr>
            </table>
            <table class="row" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 100%; display: block; padding: 0;">
              <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                <td class="wrapper last" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; position: relative; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0;" align="left" valign="top">
                  <table class="twelve columns" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; width: 580px; margin: 0 auto; padding: 0;">
                    <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                      <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; line-height: 1.3; font-size: 16px; text-align: left; margin: 0; padding: 0 20px 10px;" align="left" valign="top">
                        <p style="text-align: center; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; line-height: 1.3; margin: 0 0 10px; padding: 0;" align="center"><span style="font-size: 13px;">Если вы не оформляли заказ на нашем сайте, проигнорируйте это письмо.</span></p>
                      </td>
                    </tr>
                  </table>
                </td>
              </tr>
            </table>
            <table style="width: 100%; border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; padding: 0;">
              <tbody><tr style="vertical-align: top; text-align: left; padding: 0;" align="left"><td class="" style="font-size: 1px; line-height: 0; word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; text-align: left; width: 100%; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; background: transparent repeat center center; margin: 0; padding: 61px 0px 0px;" align="left" bgcolor="transparent" valign="top">&nbsp;</td>
              </tr></tbody></table>
              <table class="table-block" width="100%" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; padding: 0;">
                <tbody><tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                  <td class="" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; text-align: left; width: 100%; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; background: #ffffff repeat center center; margin: 0; padding: 0px 20px;" align="left" bgcolor="#ffffff" valign="top">
                  <p style="text-align: left; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; line-height: 1.3; margin: 0 0 10px; padding: 0;" align="left">С наилучшими пожеланиями,<br><span style="color: #f5b333;"><strong>MFC Foods</strong></span></p>
                  </td>
                </tr>
                </tbody></table>
              <table class="table-full" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; background: transparent repeat center center; padding: 0;" width="100%" bgcolor="transparent">
                <tbody>
                  <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                    <td height="100%" style="background-color: transparent; width: 86.31047227926078% !important; height: 100% !important; word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; text-align: left; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; margin: 0; padding: 0;" width="86.31047227926078%" align="left" bgcolor="transparent" valign="top">
                      <table class="table-block" width="100%" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; padding: 0;">
                        <tbody><tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                          <td class="" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; text-align: left; width: 100%; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; background: transparent repeat center center; margin: 0; padding: 10px 10px 0px 20px;" align="left" bgcolor="transparent" valign="top">
                          <p style="color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; text-align: left; line-height: 1.3; margin: 0 0 10px; padding: 0;" align="left"><span style="font-size: 13px;">Если у вас есть вопросы, мы будем рады вам помочь: deposite@ucg.ru</span></p>
                          </td>
                        </tr>
                        </tbody></table>
                      <table style="width: 100%; border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; padding: 0;">
                        <tbody><tr style="vertical-align: top; text-align: left; padding: 0;" align="left"><td class="" style="font-size: 1px; line-height: 0; word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; text-align: left; width: 100%; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; background: transparent repeat center center; margin: 0; padding: 28px 0px 0px;" align="left" bgcolor="transparent" valign="top">&nbsp;</td>
                        </tr></tbody></table>
                      <table width="100%" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; padding: 0;">
                        <tbody><tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                          <td style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; text-align: left; width: 100%; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; margin: 0; padding: 0px;" align="left" valign="top">
                          <hr class="" style="color: #d9d9d9; background-color: #d9d9d9; height: 1px; border: none;">
                          </td>
                        </tr>
                        </tbody></table>
                    </td>
                  </tr>
                </tbody></table>
              <table class="table-full" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; background: transparent repeat center center; padding: 0;" width="100%" bgcolor="transparent">
                <tbody>
                  <tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                    <td height="100%" style="background-color: transparent; width: 37.22221115241636% !important; height: 100% !important; word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; text-align: left; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; margin: 0; padding: 0;" width="37.22221115241636%" align="left" bgcolor="transparent" valign="top">
                      <table align="left" style="border-collapse: collapse; border-spacing: 0; overflow: hidden; width: 100%; vertical-align: top; text-align: left; padding: 0; border: 0;">
                        <tbody><tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                          <td align="left" style="text-align: left; word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; width: 100%; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; margin: 0; padding: 0px 0px 0px 20px;" valign="top">
                          <!--[if mso]>
                          <!-- --- ОБНОВЛЕНИЕ: CID ссылка для logo2 в условных комментариях Outlook --->
                          <img alt='Без заголовка.png' src='cid:logo2' width='160'>
                          <![endif]-->
                          <!--[if !mso]> <!---->
                          <!-- --- ОБНОВЛЕНИЕ: CID ссылка для logo2 --->
                          <!-- ---img alt="Без заголовка.png" class="left" height="31" src="cid:logo2" style="width: 160px !important; height: 31px; outline: none; text-decoration: none; -ms-interpolation-mode: bicubic; max-width: 100%; float: left; clear: both; display: block;" width="160" align="left" --->
                          <!-- <![endif]-->
                          </td>
                        </tr>
                        </tbody></table>
                    </td>
                    <td height="100%" style="background-color: transparent; width: 62.7777963099631% !important; height: 100% !important; word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; text-align: left; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; margin: 0; padding: 0;" width="62.7777963099631%" align="left" bgcolor="transparent" valign="top">
                      <table class="table-block" width="100%" style="border-spacing: 0; border-collapse: collapse; vertical-align: top; text-align: left; padding: 0;">
                        <tbody><tr style="vertical-align: top; text-align: left; padding: 0;" align="left">
                          <td class="" style="word-break: break-word; -webkit-hyphens: none; -moz-hyphens: none; hyphens: none; border-collapse: collapse !important; vertical-align: top; text-align: left; width: 100%; color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; background: transparent repeat center center; margin: 0; padding: 0px 10px;" align="left" bgcolor="transparent" valign="top">
                          <p style="color: #39354e; font-family: Arial, sans-serif; font-weight: normal; font-size: 16px; text-align: left; line-height: 1.3; margin: 0 0 10px; padding: 0;" align="left"><span style="line-height: 8px; font-size: 8px;">ООО "ЭМЭФСИ"&nbsp; ИНН 7714867426,&nbsp; ОГРН 1127746192716&nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; </span><span style="font-size: 8px; background-color: transparent;"><br>127018, Россия, г. Москва, ул. Двинцев, д.12, корп. 1, 1 этаж, Пом.I, ком. 30а</span></p>
                          </td>
                        </tr>
                        </tbody></table>
                    </td>
                  </tr>
                </tbody></table>
          </td>
        </tr>
      </table>
    </center>
  </td>
</tr>
</table>
</body>
</html>
"""
            # 4. Создаем объект сообщения
            # --- ОБНОВЛЕНИЕ: Используем 'related' для встраивания изображений ---
            msg = MIMEMultipart('related') 
            # Тема письма
            msg['Subject'] = f"Счёт на оплату заказа № {order_id_int}"
            # От кого (используем sender_email из конфига)
            if mail_config['sender_email']:
                 msg['From'] = f"{mail_config['sender_name']} <{mail_config['sender_email']}>"
            else:
                 # Если sender_email не задан, используем любой, например, логин системы или захардкодить
                 # msg['From'] = "noreply@yourcompany.com" # Замените на ваш реальный адрес
                 msg['From'] = "OrderSystem <ordersystem@yourcompany.com>" # Пример
            # Кому
            msg['To'] = client_email
            # 5. Добавляем HTML-часть
            msg_html = MIMEText(html_body, 'html', 'utf-8')
            msg.attach(msg_html)
            # 6. --- НОВОЕ: Добавляем локальные изображения как вложения ---
            script_dir = os.path.dirname(os.path.abspath(__file__)) # Путь к директории скрипта
            # --- Добавляем logo1.png ---
            logo1_path = os.path.join(script_dir, 'logo1.png')
            if os.path.exists(logo1_path):
                with open(logo1_path, 'rb') as f:
                    img1 = MIMEImage(f.read())
                    img1.add_header('Content-ID', '<logo1>')
                    img1.add_header('Content-Disposition', 'inline', filename='logo1.png')
                    msg.attach(img1)
                logging.debug(f"Изображение {logo1_path} добавлено как вложение CID:logo1")
            else:
                logging.warning(f"Файл изображения {logo1_path} не найден. Логотип не будет отображаться.")
            # --- Добавляем logo2.png ---
            logo2_path = os.path.join(script_dir, 'logo2.png')
            if os.path.exists(logo2_path):
                with open(logo2_path, 'rb') as f:
                    img2 = MIMEImage(f.read())
                    img2.add_header('Content-ID', '<logo2>')
                    img2.add_header('Content-Disposition', 'inline', filename='logo2.png')
                    msg.attach(img2)
                logging.debug(f"Изображение {logo2_path} добавлено как вложение CID:logo2")
            else:
                logging.warning(f"Файл изображения {logo2_path} не найден. Логотип не будет отображаться.")
            # 7. Генерируем QR-код (закомментировано, как в оригинале)
            #qr_img_buffer = BytesIO()
            # Используем URL с /q для QR-кода
            #qr = qrcode.QRCode(version=1, box_size=10, border=5)
            #qr.add_data(payment_url_qr) 
            #qr.make(fit=True)
            #img = qr.make_image(fill_color="black", back_color="white")
            #img.save(qr_img_buffer, format='PNG')
            #qr_img_buffer.seek(0)
            # 8. Добавляем QR-код как вложение (закомментировано, как в оригинале)
            #msg_image = MIMEImage(qr_img_buffer.read())
            #msg_image.add_header('Content-ID', '<qrcode>')
            #msg_image.add_header('Content-Disposition', 'inline', filename='qrcode.png')
            #msg.attach(msg_image)
            # 9. Отправляем письмо
            logging.info(f"Начинаем отправку email для заказа {order_id_int} на {client_email}...")
            server = None
            try:
                # Подключаемся к SMTP серверу
                server = smtplib.SMTP(self.config['MAIL']['smtp_server'], self.config['MAIL']['smtp_port'])
                logging.debug(f"Подключились к SMTP серверу {self.config['MAIL']['smtp_server']}:{self.config['MAIL']['smtp_port']}")
                # Включаем TLS, если требуется (редко для порта 25)
                if self.config['MAIL']['use_tls']:
                    server.starttls()
                    logging.debug("TLS включен.")
                # Логинимся ТОЛЬКО если указан пароль
                # Для Exchange 2019 через порт 25 аутентификация часто не требуется
                if self.config['MAIL']['sender_password']: # Проверяем, есть ли пароль
                    server.login(self.config['MAIL']['sender_email'], self.config['MAIL']['sender_password'])
                    logging.debug("Успешная аутентификация на SMTP сервере.")
                else:
                    logging.debug("Аутентификация не требуется или не настроена.")
                # Отправляем письмо
                # Если From не был установлен правильно, можно указать его здесь явно
                # server.sendmail("ordersystem@yourcompany.com", client_email, msg.as_string())
                server.sendmail(msg['From'], client_email, msg.as_string()) 
                logging.info(f"Email успешно отправлен для заказа {order_id_int} на {client_email}.")
                self.log_to_widget(f"📧 Письмо отправлено клиенту {client_name} ({client_email})")
                return True
            except smtplib.SMTPAuthenticationError as auth_err:
                error_msg = f"Ошибка аутентификации SMTP для {self.config['MAIL']['sender_email']}: {auth_err}"
                logging.error(error_msg)
                self.log_to_widget(f"❌ Ошибка аутентификации SMTP: {auth_err}")
                return False
            except smtplib.SMTPRecipientsRefused as recip_err:
                error_msg = f"SMTP отклонил адрес получателя {client_email}: {recip_err}"
                logging.error(error_msg)
                self.log_to_widget(f"❌ Ошибка SMTP (получатель): {recip_err}")
                return False
            except smtplib.SMTPException as smtp_err:
                error_msg = f"Ошибка SMTP при отправке на {client_email}: {smtp_err}"
                logging.error(error_msg)
                self.log_to_widget(f"❌ Ошибка SMTP: {smtp_err}")
                return False
            except Exception as e:
                error_msg = f"Неизвестная ошибка при отправке email на {client_email}: {e}"
                logging.error(error_msg, exc_info=True)
                self.log_to_widget(f"❌ Ошибка отправки email: {e}")
                return False
            finally:
                if server:
                    try:
                        server.quit()
                        logging.debug("Соединение с SMTP сервером закрыто.")
                    except:
                        pass # Игнорируем ошибки закрытия
        except Exception as e:
            error_msg = f"Критическая ошибка в send_payment_email для заказа {order_id_int}: {e}"
            logging.error(error_msg, exc_info=True)
            self.log_to_widget(f"❌ Критическая ошибка отправки email: {e}")
            return False

# --- Запуск ---
if __name__ == "__main__":
    logging.info("Запуск комбинированного приложения...")
    try:
        root = tk.Tk()
        app = CombinedApp(root)
        # Показываем сообщение об ошибке подключения после инициализации GUI
        if not app.conn:
            messagebox.showerror("❌ Критическая ошибка", "Не удалось подключиться к базе данных. Проверьте настройки в conf.ini.")
        root.mainloop()
        logging.info("Комбинированное приложение закрыто.")
    except Exception as e:
        logging.critical(f"Критическая ошибка при запуске: {e}")
        messagebox.showerror("❌ Критическая ошибка", f"Приложение не запустилось:\n{e}")
